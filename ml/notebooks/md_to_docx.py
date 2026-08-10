from docx import Document
import sys

def md_to_docx(md_path, docx_path):
    doc = Document()
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.rstrip('\n')
        if not line:
            doc.add_paragraph('')
            continue
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('- '):
            # simple bullet grouping: add as paragraph with list style
            doc.add_paragraph(line[2:].strip(), style='List Bullet')
        else:
            doc.add_paragraph(line)

    doc.save(docx_path)

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print('Usage: python md_to_docx.py input.md output.docx')
        sys.exit(1)
    md_to_docx(sys.argv[1], sys.argv[2])
